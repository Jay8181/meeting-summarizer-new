import speech_recognition as sr
import pyttsx3

# Initialize the recognizer
r = sr.Recognizer()
MyText = ''
# Function to convert text to
# speech
def SpeakText(command):
	
	# Initialize the engine
	engine = pyttsx3.init()
	engine.say(command)
	engine.runAndWait()

def ReturnTranscript():	
# Loop infinitely for user to
# speak
    while(1):
        
        # Exception handling to handle
        # exceptions at the runtime
        try:
            
            # use the microphone as source for input.
            with sr.Microphone() as source2:
                
                # wait for a second to let the recognizer
                # adjust the energy threshold based on
                # the surrounding noise level
                r.adjust_for_ambient_noise(source2, duration=0.12)
                
                #listens for the user's input
                audio2 = r.listen(source2,timeout = 120)
                
                # Using google to recognize audio
                global MyText
                MyText = r.recognize_google(audio2)
                MyText = MyText.lower()

                print("Did you say ",MyText)
                break
        except sr.RequestError as e:
            print("Could not request results; {0}".format(e))
            
        except sr.UnknownValueError:
            print("unknown error occurred")


'''
NLTK MODEL CODE
'''

# Tokenizing Sentences
from nltk.tokenize import sent_tokenize
# Tokenizing Words
from nltk.tokenize import word_tokenize
import nltk
from string import punctuation
from nltk.corpus import stopwords
nltk.download('stopwords')
nltk.download('punkt')

# Cleaning text that is got from meet transcript
def clean(text):
	sample = text.split('')
	sample.pop(0)
	clean_text = ""
	i = 0
	for t in sample:
		if i % 2 != 0:
			clean_text += str(t)
		i += 1
	return clean_text


# Finding list of stopwords ( Stopwords are
# those which do not add meaning to sentence)
stop_words = set(stopwords.words("english"))

# Tokenize
def Wtokenize(text):
	words = word_tokenize(text)
	return words


# Frequency table will be storing frequency of each word
# appearing in input text after removing stop words
# Need: It will be used for finding most relevant sentences
# as we will be applying this dictionary on every sentence
# and find its importance over other
def gen_freq_table(text):
	freqTable = dict()
	words = Wtokenize(text)
	
	for word in words:
		word = word.lower()
		if word in stop_words:
			continue
		if word in freqTable:
			freqTable[word] += 1
		else:
			freqTable[word] = 1
	return freqTable

# Sentence Tokenize
def Stokenize(text):
	sentences = sent_tokenize(text)
	return sentences

# Storing Sentence Scores
def gen_rank_sentences_table(text):

	# dictionary storing value for each sentence
	sentenceValue = dict()
	
	# Calling function gen_freq_table to get frequency
	# of words
	freqTable = gen_freq_table(text)
	
	# Calling list of sentences after tokenization
	sentences = Stokenize(text)

	for sentence in sentences:
		for word, freq in freqTable.items():
			if word in sentence.lower():
				if sentence in sentenceValue:
					sentenceValue[sentence] += freq
				else:
					sentenceValue[sentence] = freq
	return sentenceValue


def summary(text):
	sum = 0
	sentenceValue = gen_rank_sentences_table(text)
	for sentence in sentenceValue:
		sum += sentenceValue[sentence]
	avg = int(sum / (len(sentenceValue)+1))
	summary = ""
	sentences = Stokenize(text)
	for sentence in sentences:
		if (sentence in sentenceValue) and (sentenceValue[sentence] > (1.2 * avg)):
			summary += " " + sentence
	return summary


def mainFunc(inp_text, spkrr):

	summary_text = summary(inp_text)
	print(summary_text)
	global path
	my_doc = docx.Document(path+"-summary.docx")
	my_doc.add_paragraph("[" + spkrr + "]" + "\n"+summary_text)
	my_doc.save(path+"-summary.docx")

	return summary_text



import docx
timestampp = []
speakerr = []
text = []
s='' 
def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        liss =  para.text.split('\n')
        timestampp.append(liss[0])
        speakerr.append(liss[1])
        text.append(liss[2])
        global s
        sz = len(timestampp)
      
        if(sz > 1 and speakerr[sz-2] == liss[1]):
            s = s + liss[2]
        else:
            print(speakerr[sz-2])
            mainFunc(s, speakerr[sz-2])
            s = liss[2]
    fullText.append(para.text)
    return '\n'.join(fullText)


import os  
# Parent Directory path
parent_dir = "D:\\ML-summarizer\\"
lst = os.listdir(parent_dir)
number_files = len(lst)
directory = "Meeting-" + str(number_files)
  
# Path
path = os.path.join(parent_dir, directory)
#os.mkdir(path)
my_doc = docx.Document()
my_doc.add_paragraph("Summary of the Meeting: ")
my_doc.save(path+"-summary.docx")

getText("C:\\Users\\kumar\\Downloads\\Transcript_50895c2c-83d4-42d9-b556-a58ba98c9066.docx")


